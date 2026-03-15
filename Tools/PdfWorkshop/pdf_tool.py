"""
PDF Tool - Local Machine Utility
=================================
A powerful GUI tool for:
- Splitting PDF files (each page or custom groups)
- Converting PDF to Word (.docx), Text (.txt), Images (.png/.jpg)
- Batch operations and metadata viewing

Author: Generated for local use
Requirements: See requirements.txt
"""

import os
import sys
import json
import threading
import subprocess
from pathlib import Path
from typing import List, Dict, Tuple, Optional
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext

# ─────────────────────────────────────────────────────────────────────────────
# Dependency checks with user-friendly messages
# ─────────────────────────────────────────────────────────────────────────────

MISSING = []

try:
    from pypdf import PdfReader, PdfWriter
except ImportError:
    MISSING.append("pypdf")

try:
    from docx import Document
    from docx.shared import Inches, Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    MISSING.append("python-docx")

try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False
    # Not added to MISSING — optional image export, handled gracefully

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    MISSING.append("pdfplumber")


# ─────────────────────────────────────────────────────────────────────────────
# Core PDF Operations
# ─────────────────────────────────────────────────────────────────────────────

def get_page_count(pdf_path: str) -> int:
    reader = PdfReader(pdf_path)
    return len(reader.pages)


def split_each_page(pdf_path: str, output_dir: str, prefix: str = "page") -> List[str]:
    """Split every page into its own PDF file."""
    reader = PdfReader(pdf_path)
    os.makedirs(output_dir, exist_ok=True)
    outputs = []
    for i, page in enumerate(reader.pages):
        writer = PdfWriter()
        writer.add_page(page)
        out_path = os.path.join(output_dir, f"{prefix}_{i+1}.pdf")
        with open(out_path, "wb") as f:
            writer.write(f)
        outputs.append(out_path)
    return outputs


def split_custom(pdf_path: str, groups: List[List[int]], output_dir: str, prefix: str = "split") -> List[str]:
    """
    Split PDF into groups of pages.
    groups: list of lists, e.g. [[1,3], [2,5,7]] — each inner list = one output PDF
    Pages are 1-indexed.
    """
    reader = PdfReader(pdf_path)
    total = len(reader.pages)
    os.makedirs(output_dir, exist_ok=True)
    outputs = []

    for idx, group in enumerate(groups):
        writer = PdfWriter()
        valid = []
        for pg in group:
            if 1 <= pg <= total:
                writer.add_page(reader.pages[pg - 1])
                valid.append(pg)
            # silently skip out-of-range pages
        if valid:
            label = "_".join(str(p) for p in valid)
            out_path = os.path.join(output_dir, f"{prefix}_group{idx+1}_pages{label}.pdf")
            with open(out_path, "wb") as f:
                writer.write(f)
            outputs.append(out_path)
    return outputs


def split_range(pdf_path: str, start: int, end: int, output_dir: str, prefix: str = "range") -> str:
    """Extract a contiguous range of pages (1-indexed, inclusive)."""
    reader = PdfReader(pdf_path)
    total = len(reader.pages)
    start = max(1, start)
    end = min(total, end)
    os.makedirs(output_dir, exist_ok=True)

    writer = PdfWriter()
    for i in range(start - 1, end):
        writer.add_page(reader.pages[i])
    out_path = os.path.join(output_dir, f"{prefix}_pages{start}-{end}.pdf")
    with open(out_path, "wb") as f:
        writer.write(f)
    return out_path


# ─────────────────────────────────────────────────────────────────────────────
# Conversion Operations
# ─────────────────────────────────────────────────────────────────────────────

def pdf_to_text(pdf_path: str, output_path: str, pages: Optional[List[int]] = None) -> str:
    """Convert PDF to plain text using pdfplumber."""
    if not PDFPLUMBER_AVAILABLE:
        raise ImportError("pdfplumber is required for text extraction.")
    with pdfplumber.open(pdf_path) as pdf:
        all_text = []
        total = len(pdf.pages)
        target = pages if pages else list(range(1, total + 1))
        for pg_num in target:
            if 1 <= pg_num <= total:
                page = pdf.pages[pg_num - 1]
                text = page.extract_text() or ""
                all_text.append(f"{'='*50}\nPAGE {pg_num}\n{'='*50}\n{text}\n")
    content = "\n".join(all_text)
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(content)
    return output_path


def pdf_to_docx(pdf_path: str, output_path: str, pages: Optional[List[int]] = None) -> str:
    """Convert PDF to Word document preserving text structure."""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is required for Word conversion.")
    if not PDFPLUMBER_AVAILABLE:
        raise ImportError("pdfplumber is required for text extraction.")

    doc = Document()
    doc.add_heading(Path(pdf_path).stem, level=0)

    with pdfplumber.open(pdf_path) as pdf:
        total = len(pdf.pages)
        target = pages if pages else list(range(1, total + 1))
        for pg_num in target:
            if 1 <= pg_num <= total:
                page = pdf.pages[pg_num - 1]
                text = page.extract_text() or ""
                doc.add_heading(f"Page {pg_num}", level=2)
                # Split into paragraphs on blank lines
                for para in text.split("\n\n"):
                    para = para.strip()
                    if para:
                        doc.add_paragraph(para)
                doc.add_page_break()

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    return output_path


def pdf_to_images(pdf_path: str, output_dir: str, fmt: str = "png",
                   pages: Optional[List[int]] = None, dpi: int = 150) -> List[str]:
    """Convert PDF pages to images (requires pdf2image + poppler)."""
    if not PDF2IMAGE_AVAILABLE:
        raise ImportError(
            "pdf2image is required. Also install poppler:\n"
            "  Windows: https://github.com/oschwartz10612/poppler-windows\n"
            "  Mac: brew install poppler\n"
            "  Linux: sudo apt install poppler-utils"
        )
    os.makedirs(output_dir, exist_ok=True)
    kwargs = {"dpi": dpi, "fmt": fmt}
    if pages:
        kwargs["first_page"] = min(pages)
        kwargs["last_page"] = max(pages)

    images = convert_from_path(pdf_path, **kwargs)
    out_paths = []
    total = get_page_count(pdf_path)
    pg_list = pages if pages else list(range(1, total + 1))

    for i, img in enumerate(images):
        pg_num = pg_list[i] if i < len(pg_list) else i + 1
        out_path = os.path.join(output_dir, f"page_{pg_num}.{fmt}")
        img.save(out_path)
        out_paths.append(out_path)
    return out_paths


def get_metadata(pdf_path: str) -> Dict:
    """Extract PDF metadata."""
    reader = PdfReader(pdf_path)
    meta = reader.metadata or {}
    size = os.path.getsize(pdf_path)
    return {
        "File": os.path.basename(pdf_path),
        "Pages": len(reader.pages),
        "Size": f"{size / 1024:.1f} KB ({size / (1024*1024):.2f} MB)",
        "Title": meta.get("/Title", "—"),
        "Author": meta.get("/Author", "—"),
        "Subject": meta.get("/Subject", "—"),
        "Creator": meta.get("/Creator", "—"),
        "Producer": meta.get("/Producer", "—"),
    }


# ─────────────────────────────────────────────────────────────────────────────
# Parse page input like "1,3,5-8,10"
# ─────────────────────────────────────────────────────────────────────────────

def parse_pages(page_str: str, total: int) -> List[int]:
    """Parse page string like '1,3,5-8,10' into sorted unique list."""
    pages = set()
    for part in page_str.split(","):
        part = part.strip()
        if "-" in part:
            a, b = part.split("-", 1)
            try:
                a, b = int(a.strip()), int(b.strip())
                pages.update(range(a, b + 1))
            except ValueError:
                pass
        else:
            try:
                pages.add(int(part))
            except ValueError:
                pass
    return sorted(p for p in pages if 1 <= p <= total)


def parse_groups(groups_str: str, total: int) -> List[List[int]]:
    """
    Parse groups like "1,10 | 3,5 | 2-4,7"
    Each pipe-separated chunk becomes one output PDF.
    """
    groups = []
    for chunk in groups_str.split("|"):
        chunk = chunk.strip()
        if chunk:
            pages = parse_pages(chunk, total)
            if pages:
                groups.append(pages)
    return groups


# ─────────────────────────────────────────────────────────────────────────────
# GUI Application
# ─────────────────────────────────────────────────────────────────────────────

class PDFToolApp(tk.Tk):
    DARK_BG     = "#1a1a2e"
    PANEL_BG    = "#16213e"
    ACCENT      = "#0f3460"
    HIGHLIGHT   = "#e94560"
    TEXT_LIGHT  = "#eaeaea"
    TEXT_DIM    = "#8892a4"
    SUCCESS     = "#00b894"
    WARNING     = "#fdcb6e"
    FONT_HEAD   = ("Consolas", 18, "bold")
    FONT_LABEL  = ("Consolas", 10)
    FONT_BODY   = ("Consolas", 10)
    FONT_SMALL  = ("Consolas", 9)
    BTN_FONT    = ("Consolas", 10, "bold")

    def __init__(self):
        super().__init__()
        self.title("PDF Tool — Local Utility")
        self.configure(bg=self.DARK_BG)
        self.resizable(True, True)
        self.geometry("900x680")
        self.minsize(800, 600)

        self.pdf_path = tk.StringVar()
        self.output_dir = tk.StringVar()
        self.page_count = 0
        self.log_lines = []

        self._build_ui()
        self._check_deps()

    # ── UI Construction ───────────────────────────────────────────────────────

    def _build_ui(self):
        # Header
        hdr = tk.Frame(self, bg=self.DARK_BG, pady=12)
        hdr.pack(fill="x", padx=20)
        tk.Label(hdr, text="⬛ PDF TOOL", font=self.FONT_HEAD,
                 bg=self.DARK_BG, fg=self.HIGHLIGHT).pack(side="left")
        tk.Label(hdr, text="  split · convert · extract",
                 font=("Consolas", 11), bg=self.DARK_BG, fg=self.TEXT_DIM).pack(side="left", pady=4)

        # File selection row
        fs = tk.Frame(self, bg=self.PANEL_BG, padx=16, pady=10)
        fs.pack(fill="x", padx=20, pady=(0, 4))
        tk.Label(fs, text="PDF FILE:", font=self.FONT_LABEL, bg=self.PANEL_BG,
                 fg=self.TEXT_DIM, width=12, anchor="w").grid(row=0, column=0, sticky="w")
        tk.Entry(fs, textvariable=self.pdf_path, font=self.FONT_BODY,
                 bg=self.ACCENT, fg=self.TEXT_LIGHT, insertbackground=self.TEXT_LIGHT,
                 relief="flat", bd=0).grid(row=0, column=1, sticky="ew", padx=(4, 4))
        self._btn(fs, "Browse", self._browse_pdf).grid(row=0, column=2)
        self._btn(fs, "Info", self._show_info, color=self.TEXT_DIM).grid(row=0, column=3, padx=(4, 0))

        tk.Label(fs, text="OUTPUT DIR:", font=self.FONT_LABEL, bg=self.PANEL_BG,
                 fg=self.TEXT_DIM, width=12, anchor="w").grid(row=1, column=0, sticky="w", pady=(6,0))
        tk.Entry(fs, textvariable=self.output_dir, font=self.FONT_BODY,
                 bg=self.ACCENT, fg=self.TEXT_LIGHT, insertbackground=self.TEXT_LIGHT,
                 relief="flat", bd=0).grid(row=1, column=1, sticky="ew", padx=(4, 4), pady=(6,0))
        self._btn(fs, "Browse", self._browse_output).grid(row=1, column=2, pady=(6,0))
        fs.columnconfigure(1, weight=1)

        self.page_info_lbl = tk.Label(fs, text="", font=self.FONT_SMALL,
                                       bg=self.PANEL_BG, fg=self.SUCCESS)
        self.page_info_lbl.grid(row=2, column=0, columnspan=4, sticky="w", pady=(4, 0))

        # Notebook tabs
        style = ttk.Style(self)
        style.theme_use("clam")
        style.configure("TNotebook", background=self.DARK_BG, borderwidth=0)
        style.configure("TNotebook.Tab", background=self.ACCENT, foreground=self.TEXT_DIM,
                        font=self.BTN_FONT, padding=[14, 6])
        style.map("TNotebook.Tab",
                  background=[("selected", self.HIGHLIGHT)],
                  foreground=[("selected", "#ffffff")])

        nb = ttk.Notebook(self)
        nb.pack(fill="both", expand=True, padx=20, pady=6)

        self.tab_split = tk.Frame(nb, bg=self.PANEL_BG)
        self.tab_convert = tk.Frame(nb, bg=self.PANEL_BG)
        nb.add(self.tab_split,   text="  SPLIT  ")
        nb.add(self.tab_convert, text=" CONVERT ")

        self._build_split_tab()
        self._build_convert_tab()

        # Log area
        log_frame = tk.Frame(self, bg=self.DARK_BG)
        log_frame.pack(fill="both", expand=False, padx=20, pady=(0, 12))
        tk.Label(log_frame, text="LOG", font=self.FONT_SMALL,
                 bg=self.DARK_BG, fg=self.TEXT_DIM).pack(anchor="w")
        self.log_box = scrolledtext.ScrolledText(
            log_frame, height=7, font=("Consolas", 9),
            bg="#0d1117", fg=self.TEXT_LIGHT, insertbackground=self.TEXT_LIGHT,
            relief="flat", state="disabled", bd=0
        )
        self.log_box.pack(fill="both", expand=True)

    def _build_split_tab(self):
        f = self.tab_split
        pad = dict(padx=18, pady=8)

        # Mode selector
        self.split_mode = tk.StringVar(value="each")
        mode_frame = tk.Frame(f, bg=self.PANEL_BG)
        mode_frame.pack(fill="x", **pad)
        tk.Label(mode_frame, text="SPLIT MODE:", font=self.FONT_LABEL,
                 bg=self.PANEL_BG, fg=self.TEXT_DIM).pack(side="left")
        for val, label in [("each", "Each page → own PDF"),
                            ("range", "Page range → one PDF"),
                            ("custom", "Custom groups")]:
            rb = tk.Radiobutton(mode_frame, text=label, variable=self.split_mode,
                                value=val, command=self._update_split_ui,
                                bg=self.PANEL_BG, fg=self.TEXT_LIGHT,
                                selectcolor=self.HIGHLIGHT,
                                activebackground=self.PANEL_BG, font=self.FONT_BODY)
            rb.pack(side="left", padx=(12, 0))

        # Separator
        tk.Frame(f, bg=self.ACCENT, height=1).pack(fill="x", padx=18)

        # Dynamic area
        self.split_dynamic = tk.Frame(f, bg=self.PANEL_BG)
        self.split_dynamic.pack(fill="x", **pad)
        self._update_split_ui()

        # Prefix
        pf = tk.Frame(f, bg=self.PANEL_BG)
        pf.pack(fill="x", padx=18, pady=(0, 8))
        tk.Label(pf, text="FILE PREFIX:", font=self.FONT_LABEL, bg=self.PANEL_BG,
                 fg=self.TEXT_DIM, width=14, anchor="w").pack(side="left")
        self.split_prefix = tk.StringVar(value="split")
        tk.Entry(pf, textvariable=self.split_prefix, font=self.FONT_BODY,
                 bg=self.ACCENT, fg=self.TEXT_LIGHT, insertbackground=self.TEXT_LIGHT,
                 relief="flat", bd=0, width=20).pack(side="left", padx=(4, 0))

        # Run button
        self._btn(f, "▶  RUN SPLIT", self._run_split, big=True).pack(pady=12)

    def _build_convert_tab(self):
        f = self.tab_convert
        pad = dict(padx=18, pady=8)

        # Format selector
        self.conv_format = tk.StringVar(value="txt")
        fmt_frame = tk.Frame(f, bg=self.PANEL_BG)
        fmt_frame.pack(fill="x", **pad)
        tk.Label(fmt_frame, text="OUTPUT FORMAT:", font=self.FONT_LABEL,
                 bg=self.PANEL_BG, fg=self.TEXT_DIM).pack(side="left")
        for val, label in [("txt", "Plain Text (.txt)"),
                            ("docx", "Word Document (.docx)"),
                            ("png", "Images PNG (.png)"),
                            ("jpg", "Images JPG (.jpg)")]:
            rb = tk.Radiobutton(fmt_frame, text=label, variable=self.conv_format,
                                value=val, command=self._update_conv_ui,
                                bg=self.PANEL_BG, fg=self.TEXT_LIGHT,
                                selectcolor=self.HIGHLIGHT,
                                activebackground=self.PANEL_BG, font=self.FONT_BODY)
            rb.pack(side="left", padx=(12, 0))

        tk.Frame(f, bg=self.ACCENT, height=1).pack(fill="x", padx=18)

        # Page filter
        pg_frame = tk.Frame(f, bg=self.PANEL_BG)
        pg_frame.pack(fill="x", **pad)
        tk.Label(pg_frame, text="PAGES (optional):", font=self.FONT_LABEL,
                 bg=self.PANEL_BG, fg=self.TEXT_DIM, width=20, anchor="w").pack(side="left")
        self.conv_pages = tk.StringVar()
        tk.Entry(pg_frame, textvariable=self.conv_pages, font=self.FONT_BODY,
                 bg=self.ACCENT, fg=self.TEXT_LIGHT, insertbackground=self.TEXT_LIGHT,
                 relief="flat", bd=0, width=30).pack(side="left", padx=(4, 0))
        tk.Label(pg_frame, text='e.g. "1,3,5-8" or blank for all',
                 font=self.FONT_SMALL, bg=self.PANEL_BG, fg=self.TEXT_DIM).pack(side="left", padx=(8, 0))

        # DPI (for image modes)
        self.dpi_frame = tk.Frame(f, bg=self.PANEL_BG)
        self.dpi_frame.pack(fill="x", padx=18, pady=(0, 8))
        tk.Label(self.dpi_frame, text="DPI (image quality):", font=self.FONT_LABEL,
                 bg=self.PANEL_BG, fg=self.TEXT_DIM, width=20, anchor="w").pack(side="left")
        self.conv_dpi = tk.StringVar(value="150")
        tk.Entry(self.dpi_frame, textvariable=self.conv_dpi, font=self.FONT_BODY,
                 bg=self.ACCENT, fg=self.TEXT_LIGHT, insertbackground=self.TEXT_LIGHT,
                 relief="flat", bd=0, width=8).pack(side="left", padx=(4, 0))
        tk.Label(self.dpi_frame, text="(72=draft  150=normal  300=print)",
                 font=self.FONT_SMALL, bg=self.PANEL_BG, fg=self.TEXT_DIM).pack(side="left", padx=(8, 0))

        # Dynamic note
        self.conv_note = tk.Label(f, text="", font=self.FONT_SMALL,
                                   bg=self.PANEL_BG, fg=self.WARNING, wraplength=700, justify="left")
        self.conv_note.pack(fill="x", padx=18, pady=(0, 4))
        self._update_conv_ui()

        # Run button
        self._btn(f, "▶  RUN CONVERT", self._run_convert, big=True).pack(pady=12)

    # ── Dynamic UI helpers ────────────────────────────────────────────────────

    def _update_split_ui(self):
        for w in self.split_dynamic.winfo_children():
            w.destroy()
        mode = self.split_mode.get()
        f = self.split_dynamic

        if mode == "each":
            tk.Label(f, text="Every page will be saved as its own PDF file.",
                     font=self.FONT_BODY, bg=self.PANEL_BG, fg=self.TEXT_LIGHT).pack(anchor="w")

        elif mode == "range":
            tk.Label(f, text="Extract a range of pages into one PDF.",
                     font=self.FONT_SMALL, bg=self.PANEL_BG, fg=self.TEXT_DIM).pack(anchor="w")
            row = tk.Frame(f, bg=self.PANEL_BG)
            row.pack(anchor="w", pady=(4, 0))
            tk.Label(row, text="Start page:", font=self.FONT_LABEL, bg=self.PANEL_BG,
                     fg=self.TEXT_DIM).pack(side="left")
            self.range_start = tk.StringVar(value="1")
            tk.Entry(row, textvariable=self.range_start, font=self.FONT_BODY,
                     bg=self.ACCENT, fg=self.TEXT_LIGHT, insertbackground=self.TEXT_LIGHT,
                     relief="flat", bd=0, width=6).pack(side="left", padx=(4, 16))
            tk.Label(row, text="End page:", font=self.FONT_LABEL, bg=self.PANEL_BG,
                     fg=self.TEXT_DIM).pack(side="left")
            self.range_end = tk.StringVar(value="5")
            tk.Entry(row, textvariable=self.range_end, font=self.FONT_BODY,
                     bg=self.ACCENT, fg=self.TEXT_LIGHT, insertbackground=self.TEXT_LIGHT,
                     relief="flat", bd=0, width=6).pack(side="left", padx=(4, 0))

        elif mode == "custom":
            info = ('Enter groups separated by  |  — pages inside each group separated by commas.\n'
                    'Example:  1,10 | 3,5 | 2-4,7\n'
                    '→ Group 1: pages 1 & 10   →  Group 2: pages 3 & 5   →  Group 3: pages 2,3,4,7')
            tk.Label(f, text=info, font=self.FONT_SMALL, bg=self.PANEL_BG,
                     fg=self.TEXT_DIM, justify="left").pack(anchor="w")
            self.custom_groups = tk.StringVar(value="1,2 | 3,4 | 5-7")
            tk.Entry(f, textvariable=self.custom_groups, font=self.FONT_BODY,
                     bg=self.ACCENT, fg=self.TEXT_LIGHT, insertbackground=self.TEXT_LIGHT,
                     relief="flat", bd=0).pack(fill="x", pady=(6, 0))

    def _update_conv_ui(self):
        fmt = self.conv_format.get()
        img_mode = fmt in ("png", "jpg")
        if img_mode:
            self.dpi_frame.pack(fill="x", padx=18, pady=(0, 8))
        else:
            self.dpi_frame.pack_forget()

        notes = {
            "txt": "",
            "docx": "Requires python-docx. Extracts text with page headings." if DOCX_AVAILABLE else
                    "⚠  python-docx not installed — run:  pip install python-docx",
            "png": "Requires pdf2image + poppler. See README for poppler install." if PDF2IMAGE_AVAILABLE else
                   "⚠  pdf2image not installed — run:  pip install pdf2image\n"
                   "   Also install poppler (see README).",
            "jpg": "Requires pdf2image + poppler." if PDF2IMAGE_AVAILABLE else
                   "⚠  pdf2image not installed — run:  pip install pdf2image",
        }
        self.conv_note.config(text=notes.get(fmt, ""))

    # ── Actions ───────────────────────────────────────────────────────────────

    def _browse_pdf(self):
        path = filedialog.askopenfilename(
            title="Select PDF file",
            filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")]
        )
        if path:
            self.pdf_path.set(path)
            if not self.output_dir.get():
                self.output_dir.set(str(Path(path).parent / "pdf_tool_output"))
            try:
                self.page_count = get_page_count(path)
                self.page_info_lbl.config(
                    text=f"✔  {self.page_count} pages  ·  {Path(path).name}",
                    fg=self.SUCCESS
                )
            except Exception as e:
                self.page_info_lbl.config(text=f"⚠  Could not read PDF: {e}", fg=self.WARNING)
                self.page_count = 0

    def _browse_output(self):
        path = filedialog.askdirectory(title="Select output folder")
        if path:
            self.output_dir.set(path)

    def _show_info(self):
        pdf = self.pdf_path.get()
        if not pdf or not os.path.isfile(pdf):
            messagebox.showwarning("No PDF", "Please select a valid PDF file first.")
            return
        try:
            meta = get_metadata(pdf)
            msg = "\n".join(f"{k:12s}: {v}" for k, v in meta.items())
            messagebox.showinfo("PDF Metadata", msg)
        except Exception as e:
            messagebox.showerror("Error", str(e))

    def _run_split(self):
        if not self._validate():
            return
        mode = self.split_mode.get()
        pdf = self.pdf_path.get()
        out = self.output_dir.get()
        prefix = self.split_prefix.get().strip() or "split"

        def task():
            try:
                if mode == "each":
                    self._log("Splitting every page into individual PDFs …")
                    files = split_each_page(pdf, out, prefix)
                    self._log(f"✔  Done! {len(files)} files → {out}", ok=True)

                elif mode == "range":
                    try:
                        s = int(self.range_start.get())
                        e = int(self.range_end.get())
                    except ValueError:
                        self._log("⚠  Invalid range — enter integers.", err=True)
                        return
                    self._log(f"Extracting pages {s}–{e} …")
                    out_file = split_range(pdf, s, e, out, prefix)
                    self._log(f"✔  Saved: {out_file}", ok=True)

                elif mode == "custom":
                    raw = self.custom_groups.get()
                    groups = parse_groups(raw, self.page_count)
                    if not groups:
                        self._log("⚠  No valid groups found.", err=True)
                        return
                    self._log(f"Splitting into {len(groups)} custom groups …")
                    files = split_custom(pdf, groups, out, prefix)
                    for f in files:
                        self._log(f"  → {os.path.basename(f)}")
                    self._log(f"✔  Done! {len(files)} files saved.", ok=True)

                self._open_folder(out)
            except Exception as ex:
                self._log(f"ERROR: {ex}", err=True)

        threading.Thread(target=task, daemon=True).start()

    def _run_convert(self):
        if not self._validate():
            return
        pdf = self.pdf_path.get()
        out = self.output_dir.get()
        fmt = self.conv_format.get()
        pg_str = self.conv_pages.get().strip()
        pages = parse_pages(pg_str, self.page_count) if pg_str else None
        stem = Path(pdf).stem

        def task():
            try:
                if fmt == "txt":
                    self._log("Converting to plain text …")
                    out_path = os.path.join(out, stem + ".txt")
                    pdf_to_text(pdf, out_path, pages)
                    self._log(f"✔  Saved: {out_path}", ok=True)

                elif fmt == "docx":
                    self._log("Converting to Word document …")
                    out_path = os.path.join(out, stem + ".docx")
                    pdf_to_docx(pdf, out_path, pages)
                    self._log(f"✔  Saved: {out_path}", ok=True)

                elif fmt in ("png", "jpg"):
                    try:
                        dpi = int(self.conv_dpi.get())
                    except ValueError:
                        dpi = 150
                    img_dir = os.path.join(out, stem + "_images")
                    self._log(f"Converting to {fmt.upper()} images (DPI={dpi}) …")
                    files = pdf_to_images(pdf, img_dir, fmt, pages, dpi)
                    self._log(f"✔  {len(files)} images saved to: {img_dir}", ok=True)
                    out = img_dir

                self._open_folder(out)
            except ImportError as ex:
                self._log(f"⚠  Missing dependency: {ex}", err=True)
            except Exception as ex:
                self._log(f"ERROR: {ex}", err=True)

        threading.Thread(target=task, daemon=True).start()

    # ── Helpers ───────────────────────────────────────────────────────────────

    def _validate(self) -> bool:
        pdf = self.pdf_path.get()
        if not pdf or not os.path.isfile(pdf):
            messagebox.showwarning("No PDF", "Please select a valid PDF file.")
            return False
        if not self.output_dir.get():
            messagebox.showwarning("No Output", "Please select an output directory.")
            return False
        if self.page_count == 0:
            try:
                self.page_count = get_page_count(pdf)
            except Exception:
                messagebox.showerror("PDF Error", "Could not read the selected PDF.")
                return False
        return True

    def _log(self, msg: str, ok: bool = False, err: bool = False):
        def _do():
            self.log_box.config(state="normal")
            prefix = "✔  " if ok else ("✖  " if err else "   ")
            self.log_box.insert("end", prefix + msg + "\n")
            self.log_box.see("end")
            self.log_box.config(state="disabled")
        self.after(0, _do)

    def _open_folder(self, folder: str):
        """Open the output folder in the OS file explorer."""
        try:
            if sys.platform == "win32":
                os.startfile(folder)
            elif sys.platform == "darwin":
                subprocess.run(["open", folder])
            else:
                subprocess.run(["xdg-open", folder])
        except Exception:
            pass

    def _check_deps(self):
        if MISSING:
            msg = ("Some features require additional packages.\n\n"
                   "Install with:\n  pip install " + " ".join(MISSING))
            self._log("⚠  Missing packages: " + ", ".join(MISSING) + "  (see log for details)", err=True)
            self._log("Run:  pip install " + " ".join(MISSING))

    def _btn(self, parent, text, cmd, color=None, big=False):
        bg = color or self.HIGHLIGHT
        font = ("Consolas", 11, "bold") if big else self.BTN_FONT
        return tk.Button(
            parent, text=text, command=cmd,
            bg=bg, fg="#ffffff", activebackground=self.ACCENT,
            relief="flat", cursor="hand2", font=font,
            padx=14 if big else 10, pady=6 if big else 4,
            bd=0
        )


# ─────────────────────────────────────────────────────────────────────────────
# Entry point
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    if MISSING and "pypdf" in MISSING:
        print("CRITICAL: pypdf is required.  Run:  pip install pypdf")
        sys.exit(1)
    app = PDFToolApp()
    app.mainloop()
